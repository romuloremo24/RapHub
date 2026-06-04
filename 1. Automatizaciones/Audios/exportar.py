"""
Lee el archivo formateado, aplica limpieza manual de errores conocidos,
agrega sección de consideraciones y exporta a TXT + DOCX.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

TRANSCRIPCIONES = Path(__file__).parent / "Transcripciones"

# ---------------------------------------------------------------------------
# Limpieza de errores conocidos
# ---------------------------------------------------------------------------

REEMPLAZOS = [
    # Artefacto de audio — notificación o ruido de fondo
    (
        "ENTREVISTADORA: Más de un millón de personas han sido asesinadas por el gobierno de la República Dominicana.",
        "[ARTEFACTO DE AUDIO — fragmento ininteligible]"
    ),
    # Prompt de Whisper que se coló en la transcripción
    (
        "Mantener modismos, coloquialismos y palabras tal cual se dicen. Mantener modismos, coloquialismos y palabras tal cual se dicen.",
        ""
    ),
    # Error de atribución en bloque de denuncia (líneas ~175-183)
    (
        "ENTREVISTADORA: Yo esperé como cuatro días en contarle a mi hermana.\n\nNACHA: Primero le conté a una amiga, a la María. Pero hiciste la prueba de... porque son... o sea, cuando tú ibas a denunciar, te hicieron inmediatamente la constatación de lesiones.\n\nENTREVISTADORA: Yo cuando fui a denunciar, en verdad fue así.\n\nNACHA: Yo le conté a mi hermana después de como cuatro días, que había pasado. Yo ya estaba en Santiago, pero ya estaba para la caga. Me estaba ganando y yo sabía que si no buscaba a Iba, yo me iba a terminar matando. Yo lo único que quería era morirme en ese minuto. Y le dije a mi hermana, y ahí mi hermana, bueno, no andaba, me ayudó Caleta. Caleta, Caleta, Caleta, Caleta, maldito, maldito, maldito, y yo le lloraba. Le lloraba, dije, no sé cómo me había pasado esto.\n\nENTREVISTADORA: Y ahí fuimos a la clínica al tiro.",
        "NACHA: Yo esperé como cuatro días en contarle a mi hermana. Primero le conté a una amiga, a la María.\n\nENTREVISTADORA: Pero hiciste la prueba de... porque son... o sea, cuando tú ibas a denunciar, te hicieron inmediatamente la constatación de lesiones.\n\nNACHA: Yo cuando fui a denunciar, en verdad fue así. Yo le conté a mi hermana después de como cuatro días, que había pasado. Yo ya estaba en Santiago, pero ya estaba para la caga. Me estaba ganando y yo sabía que si no buscaba a Iba, yo me iba a terminar matando. Yo lo único que quería era morirme en ese minuto. Y le dije a mi hermana, y ahí mi hermana, bueno, no andaba, me ayudó Caleta. Caleta, Caleta, Caleta, Caleta, maldito, maldito, maldito, y yo le lloraba. Le lloraba, dije, no sé cómo me había pasado esto. Y ahí fuimos a la clínica al tiro."
    ),
    # Repetición de pregunta/respuesta en parte 3
    (
        "ENTREVISTADORA: ¿Y tú sabías que alguna vez se tiró como a una niña? \n\nNACHA: Nunca, jamás se la iba. De hecho esto para mí fue así como... Yo era menor de él. Es que es eso, ¿cachai? Es una cuestión como... Él tenía entonces 28, creo. 28. Entonces es una cuestión que no, como que, no sé, es como que te lo juro que me hablaron de otra persona. Por eso es que estoy tan choqueada y me ha costado un poco asimilarlo porque es como rebobinar y cambiar completamente todos los pensamientos que tenía. \n\nENTREVISTADORA: ¿Y tú sabías que alguna vez se tiró como a una niña? \n\nNACHA: Nunca, jamás se la iba. De hecho esto para mí fue así como... Yo era menor de él. Es que es eso, ¿cachai? Es una cuestión como... Él tenía entonces 28, creo. 28. Entonces es una cuestión que no, como que, no sé, es como que te lo juro que me hablaron de otra persona. Por eso es que estoy tan choqueada y me ha costado un poco asimilarlo porque es como rebobinar y cambiar completamente todos los pensamientos que tenía. Entonces es como... ¿Cómo lo asocio? Es como si viviera yo en una mentira. ",
        "ENTREVISTADORA: ¿Y tú sabías que alguna vez se tiró como a una niña?\n\nNACHA: Nunca, jamás se la iba. De hecho esto para mí fue así como... Yo era menor de él. Es que es eso, ¿cachai? Es una cuestión como... Él tenía entonces 28, creo. 28. Entonces es una cuestión que no, como que, no sé, es como que te lo juro que me hablaron de otra persona. Por eso es que estoy tan choqueada y me ha costado un poco asimilarlo porque es como rebobinar y cambiar completamente todos los pensamientos que tenía. Entonces es como... ¿Cómo lo asocio? Es como si viviera yo en una mentira."
    ),
]

CONSIDERACIONES = """
════════════════════════════════════════════════════════════
CONSIDERACIONES TÉCNICAS Y DE CONTENIDO
════════════════════════════════════════════════════════════

1. MÉTODO DE TRANSCRIPCIÓN
   Transcripción automática generada con OpenAI Whisper large-v3 vía API Groq.
   Idioma: español. Formato de audio original: MP4 (12 MB).

2. IDENTIFICACIÓN DE HABLANTES
   Se identificaron 2 participantes:
   - NACHA: relata en primera persona una experiencia de agresión sexual.
   - ENTREVISTADORA: conduce la conversación con preguntas, llamada "Sophie"
     en al menos una ocasión por Nacha.
   La atribución de turnos cortos fue revisada y corregida manualmente
   donde el modelo automático cometió errores.

3. FRAGMENTOS CON AUDIO POCO CLARO
   Las siguientes secciones presentan transcripción incompleta o con errores
   de reconocimiento (audio distorsionado, habla superpuesta o muy baja):

   a) Segmento "baño" — frase cortada:
      "Me dice como pero ven qu est llorando Le digo no es que soy al le dije
      una vez as Me dice qu Y yo le dije no al copete Y me hace muy mal y me
      siento muy mal No s por qu le dije eso pero estaba cagada y miedo no s
      enfrentarlo confrontarlo que me hiciera algo m all"
      → Audio original presenta bajo nivel o habla acelerada en este tramo.

   b) Segmento "colarse a carretes" — frase sin terminar:
      "La electr que fue en marzo Arena Santiago se llamaba y ah como que yo
      me colaba como a sus carretes"
      → Probable nombre propio o expresión chilena no reconocida por el modelo.

   c) Segmento "hongo / bosque":
      "Porque igual siempre creo que no s d pas Yo creo que fueron todos los
      que estaban en que me hicieron el hongo"
      → Posible referencia a drogas. Audio poco claro en este tramo.

4. ARTEFACTOS DETECTADOS Y ELIMINADOS
   - Una frase incoherente sobre "República Dominicana" fue detectada y
     marcada como artefacto de audio (probablemente ruido de fondo o
     notificación durante la grabación).
   - El prompt de instrucción de Whisper ("Mantener modismos, coloquialismos
     y palabras tal cual se dicen") apareció duplicado en la transcripción
     cruda y fue eliminado.

5. CONTENIDO SENSIBLE
   Esta transcripción contiene el relato de una agresión sexual. La persona
   identificada como Nacha menciona: estado de inconsciencia, imposibilidad
   de moverse, dolor físico, y pensamientos suicidas posteriores al evento.
   Manejar con confidencialidad.
"""


def limpiar(texto: str) -> str:
    for buscar, reemplazar in REEMPLAZOS:
        texto = texto.replace(buscar, reemplazar)
    # Eliminar líneas en blanco múltiples
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    # Eliminar líneas que solo tienen un espacio
    texto = re.sub(r"\n \n", "\n\n", texto)
    return texto.strip()


# ---------------------------------------------------------------------------
# Exportar TXT
# ---------------------------------------------------------------------------

def exportar_txt(texto: str, ruta: Path):
    ruta.write_text(texto, encoding="utf-8")
    print(f"  TXT -> {ruta}")


# ---------------------------------------------------------------------------
# Exportar DOCX
# ---------------------------------------------------------------------------

COLORES = {
    "NACHA": RGBColor(0x15, 0x13, 0x33),       # azul oscuro
    "ENTREVISTADORA": RGBColor(0x2E, 0x67, 0xB0),  # azul medio
    "ARTEFACTO": RGBColor(0x99, 0x99, 0x99),    # gris
    "HEADER": RGBColor(0x75, 0xC2, 0xF5),       # azul claro
}

def set_run_style(run, bold=False, color=None, size=11):
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color

def exportar_docx(texto: str, ruta: Path):
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Estilo base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lineas = texto.split("\n")
    i = 0
    en_consideraciones = False

    while i < len(lineas):
        linea = lineas[i].strip()

        # Encabezado de participantes
        if linea.startswith("PARTICIPANTES:"):
            p = doc.add_paragraph()
            r = p.add_run(linea)
            set_run_style(r, bold=True, color=COLORES["HEADER"], size=12)
            i += 1
            continue

        # Sub-líneas del encabezado
        if linea.startswith("- NACHA:") or linea.startswith("- ENTREVISTADORA:"):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.3)
            r = p.add_run(linea[2:])
            set_run_style(r, size=10, color=COLORES["HEADER"])
            i += 1
            continue

        # Separador
        if set(linea) <= {"─", "═", "=", "-"} and len(linea) > 5:
            doc.add_paragraph()
            i += 1
            continue

        # Título CONSIDERACIONES
        if "CONSIDERACIONES TÉCNICAS" in linea:
            doc.add_page_break()
            en_consideraciones = True
            p = doc.add_heading("Consideraciones técnicas y de contenido", level=1)
            i += 1
            continue

        # Subtítulos numerados dentro de consideraciones
        if en_consideraciones and re.match(r"^\d\.", linea):
            p = doc.add_heading(linea, level=2)
            i += 1
            continue

        # Artefacto
        if "[ARTEFACTO" in linea:
            p = doc.add_paragraph()
            r = p.add_run(linea)
            set_run_style(r, color=COLORES["ARTEFACTO"], size=10)
            r.italic = True
            i += 1
            continue

        # Turno de NACHA
        if linea.startswith("NACHA:"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            r_label = p.add_run("NACHA: ")
            set_run_style(r_label, bold=True, color=COLORES["NACHA"])
            r_texto = p.add_run(linea[6:].strip())
            set_run_style(r_texto, color=COLORES["NACHA"])
            i += 1
            continue

        # Turno de ENTREVISTADORA
        if linea.startswith("ENTREVISTADORA:"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            r_label = p.add_run("ENTREVISTADORA: ")
            set_run_style(r_label, bold=True, color=COLORES["ENTREVISTADORA"])
            r_texto = p.add_run(linea[15:].strip())
            set_run_style(r_texto, color=COLORES["ENTREVISTADORA"])
            i += 1
            continue

        # Líneas de consideraciones (texto normal)
        if linea:
            p = doc.add_paragraph(linea)
            p.paragraph_format.left_indent = Inches(0.3) if linea.startswith("→") or linea.startswith("a)") or linea.startswith("b)") or linea.startswith("c)") else Inches(0)

        i += 1

    doc.save(ruta)
    print(f"  DOCX -> {ruta}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    fuente = TRANSCRIPCIONES / "audio victima - esposa_formateado.txt"
    texto_raw = fuente.read_text(encoding="utf-8")

    texto_limpio = limpiar(texto_raw) + "\n\n" + CONSIDERACIONES.strip()

    nombre_base = "audio victima - esposa_final"
    exportar_txt(texto_limpio, TRANSCRIPCIONES / f"{nombre_base}.txt")
    exportar_docx(texto_limpio, TRANSCRIPCIONES / f"{nombre_base}.docx")
    print("Listo.")

if __name__ == "__main__":
    main()
